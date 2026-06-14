"""Resume file parsing: text extraction (PDF/DOCX/TXT) + structured parse.

Structured parsing uses Claude when available, with a regex-based fallback so
the system still functions without an API key. Also computes the Resume
Strength Score (spec section 9.1).
"""
import logging
import re
import time
from pathlib import Path

from . import config, db, llm

log = logging.getLogger("resume")

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        import pdfplumber
        text = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text() or "")
        return "\n".join(text).strip()
    if ext == ".docx":
        import docx
        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()
    if ext == ".txt":
        return path.read_text(encoding="utf-8", errors="replace").strip()
    raise ValueError(f"Unsupported file type: {ext}")


PARSE_SCHEMA = llm.obj_schema({
    "full_name": llm.STR,
    "email": llm.STR,
    "linkedin_url": llm.STR,
    "portfolio_url": llm.STR,
    "city": llm.STR,
    "state": llm.STR,
    "skills": llm.STR_ARR,
    "experience_level": {"type": "string", "enum": ["Entry", "Mid", "Senior", "Executive"]},
    "work_history": {
        "type": "array",
        "items": llm.obj_schema({
            "title": llm.STR, "company": llm.STR, "duration": llm.STR, "description": llm.STR,
        }),
    },
})


def parse_resume_text(raw_text: str, user_id=None) -> dict:
    """Returns dict with keys: full_name, email, linkedin_url, portfolio_url,
    city, state, skills (list), experience_level, work_history (list)."""
    text = raw_text[:config.RESUME_TEXT_LIMIT]
    result = llm.complete_json(
        "You are an expert resume parser. Extract structured data from the resume text below.\n"
        "- email: the candidate's email address (empty string if none found)\n"
        "- linkedin_url: full LinkedIn profile URL if present (also detect mentions like "
        "'LinkedIn: johndoe' and convert to https://linkedin.com/in/johndoe); empty string if none\n"
        "- portfolio_url: personal website/portfolio URL if present; empty string if none\n"
        "- city / state: the candidate's location (empty strings if not found)\n"
        "- skills: a comprehensive list of skills (technical and transferable)\n"
        "- experience_level: one of Entry, Mid, Senior, Executive\n"
        "- work_history: past positions, most recent first\n\n"
        f"Resume text:\n{text}",
        process_type="resume_parsing",
        user_id=user_id,
        schema=PARSE_SCHEMA,
        max_tokens=4096,
    )
    if result is None:
        result = _heuristic_parse(text)
    # normalise
    result["email"] = (result.get("email") or "").strip().lower()
    if not _valid_email(result["email"]):
        # one more chance with regex on the raw text
        match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", raw_text)
        result["email"] = match.group(0).lower() if match else ""
    result.setdefault("skills", [])
    result.setdefault("work_history", [])
    result["experience_level"] = result.get("experience_level") or "Mid"
    return result


def _valid_email(email: str) -> bool:
    return bool(re.fullmatch(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", email or ""))


_SKILL_HINTS = [
    "python", "java", "javascript", "typescript", "sql", "excel", "word", "powerpoint",
    "salesforce", "aws", "azure", "docker", "kubernetes", "react", "node", "c++", "c#",
    "customer service", "project management", "leadership", "scheduling", "budgeting",
    "forklift", "inventory", "logistics", "marketing", "sales", "accounting", "quickbooks",
    "data analysis", "machine learning", "communication", "training", "negotiation",
    "html", "css", "git", "linux", "tableau", "power bi", "spanish", "bilingual",
]


def _heuristic_parse(text: str) -> dict:
    lower = text.lower()
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    linkedin = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-.%]+", text, re.I)
    # first non-empty line that isn't an email/phone is usually the name
    name = ""
    for line in text.splitlines():
        line = line.strip()
        if line and "@" not in line and not re.search(r"\d{3}", line) and len(line) < 60:
            name = line
            break
    city, state = "", ""
    loc = re.search(r"([A-Z][A-Za-z .]+),\s*([A-Z]{2})\b", text)
    if loc:
        city, state = loc.group(1).strip(), loc.group(2)
    skills = sorted({s.title() for s in _SKILL_HINTS if s in lower})
    years = re.findall(r"(\d{1,2})\+?\s*years", lower)
    longest = max((int(y) for y in years), default=3)
    level = "Entry" if longest < 2 else "Mid" if longest < 7 else "Senior" if longest < 15 else "Executive"
    return {
        "full_name": name,
        "email": email_match.group(0).lower() if email_match else "",
        "linkedin_url": ("https://" + linkedin.group(0).lstrip("htps:/")) if linkedin else "",
        "portfolio_url": "",
        "city": city,
        "state": state,
        "skills": skills,
        "experience_level": level,
        "work_history": [],
    }


STRENGTH_SCHEMA = llm.obj_schema({
    "score": llm.INT,
    "summary": llm.STR,
    "top_improvements": llm.STR_ARR,
})


def compute_strength_score(raw_text: str, user_id=None) -> dict:
    """Resume Strength Score 0-100 (spec 9.1)."""
    result = llm.complete_json(
        "You are a professional resume reviewer. Score this resume on a scale of 0-100 based on:\n"
        "- Formatting and readability (10 points)\n"
        "- Keyword optimization for job searches (20 points)\n"
        "- Quantified achievements vs. vague descriptions (20 points)\n"
        "- Skills presentation and relevance (20 points)\n"
        "- Work history completeness and progression (15 points)\n"
        "- Overall professionalism (15 points)\n\n"
        f"Resume text:\n{raw_text[:config.RESUME_TEXT_LIMIT]}\n\n"
        "Return score, a one sentence summary of strengths and weaknesses, and the top 3 improvements.",
        process_type="resume_advice",
        user_id=user_id,
        schema=STRENGTH_SCHEMA,
        max_tokens=1024,
    )
    if result is None:
        # heuristic: length, numbers, sections
        score = 40
        if len(raw_text) > 1200:
            score += 15
        if len(re.findall(r"\d+%|\$\d|\d+ (?:people|projects|clients)", raw_text)) >= 3:
            score += 20
        if re.search(r"(?i)\b(experience|education|skills)\b", raw_text):
            score += 15
        result = {
            "score": min(score, 90),
            "summary": "Automated heuristic score (connect an Anthropic API key for a full AI review).",
            "top_improvements": ["Quantify achievements with numbers",
                                 "Tailor keywords to target roles",
                                 "Keep formatting simple and consistent"],
        }
    result["score"] = max(0, min(100, int(result.get("score") or 0)))
    return result


def save_resume_file(user_id, original_filename: str, content: bytes) -> Path:
    ext = Path(original_filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError("Only PDF and DOCX files are supported.")
    fname = f"{user_id or 'new'}_{int(time.time() * 1000)}{ext}"
    path = config.RESUME_DIR / fname
    path.write_bytes(content)
    return path


def create_or_update_user_from_resume(path: Path, raw_text: str, parsed: dict,
                                      label: str = "Primary Resume",
                                      manual_email: str = None) -> dict:
    """Admin upload pipeline: create the user (or update existing by email),
    attach the resume, compute strength score. Returns {user_id, created,
    needs_email, parsed, strength}."""
    email = (manual_email or parsed.get("email") or "").strip().lower()
    if not _valid_email(email):
        return {"user_id": None, "created": False, "needs_email": True,
                "parsed": parsed, "strength": None}

    existing = db.query_one("SELECT id FROM users WHERE email = ?", (email,))
    skills_json = db.jdumps(parsed.get("skills") or [])
    work_json = db.jdumps(parsed.get("work_history") or [])
    if existing:
        user_id = existing["id"]
        db.execute(
            "UPDATE users SET full_name = COALESCE(NULLIF(?, ''), full_name), "
            "city = COALESCE(NULLIF(?, ''), city), state = COALESCE(NULLIF(?, ''), state), "
            "skills = ?, experience_level = ?, work_history = ?, resume_file_path = ?, "
            "resume_raw_text = ?, linkedin_url = COALESCE(NULLIF(?, ''), linkedin_url), "
            "portfolio_url = COALESCE(NULLIF(?, ''), portfolio_url), "
            "resume_advice_pending = 1, updated_at = ? WHERE id = ?",
            (parsed.get("full_name") or "", parsed.get("city") or "", parsed.get("state") or "",
             skills_json, parsed.get("experience_level"), work_json, str(path), raw_text,
             parsed.get("linkedin_url") or "", parsed.get("portfolio_url") or "", db.now(), user_id),
        )
        created = False
    else:
        user_id = db.execute(
            "INSERT INTO users (full_name, email, city, state, skills, experience_level, "
            "work_history, resume_file_path, resume_raw_text, linkedin_url, portfolio_url, "
            "resume_advice_pending) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 1)",
            (parsed.get("full_name") or "", email, parsed.get("city") or "",
             parsed.get("state") or "", skills_json, parsed.get("experience_level"),
             work_json, str(path), raw_text, parsed.get("linkedin_url") or "",
             parsed.get("portfolio_url") or ""),
        )
        created = True

    # primary resume row
    db.execute("UPDATE user_resumes SET is_primary = 0 WHERE user_id = ?", (user_id,))
    db.execute(
        "INSERT INTO user_resumes (user_id, label, file_path, raw_text, skills, experience_level, is_primary) "
        "VALUES (?, ?, ?, ?, ?, ?, 1)",
        (user_id, label, str(path), raw_text, skills_json, parsed.get("experience_level")),
    )
    strength = compute_strength_score(raw_text, user_id)
    db.execute(
        "UPDATE users SET resume_strength_score = ?, resume_strength_summary = ? WHERE id = ?",
        (strength["score"], strength.get("summary", ""), user_id),
    )
    return {"user_id": user_id, "created": created, "needs_email": False,
            "parsed": parsed, "strength": strength}
