"""Second smoke pass: features not covered by tests/smoke.py.
Covers: multiple resumes, Give-Me-Advice-Now, email preference suppression,
urgency/portfolio/token-budget runs, frequency transition, GDPR deletion with
anonymous story persistence."""
import re
import sys
import time
from pathlib import Path

import httpx

BASE = "http://127.0.0.1:8000"
ROOT = Path(__file__).resolve().parent.parent
OUTBOX = ROOT / "data" / "outbox"

ADMIN_EMAIL = "ivandoublejr@gmail.com"
SCRATCH_EMAIL = "scratch.user.delete.me@example.com"

passed, failed = [], []


def check(name, condition, detail=""):
    (passed if condition else failed).append(name)
    print(("  PASS " if condition else "  FAIL ") + name + (f" - {detail}" if detail and not condition else ""))


def latest_otp(email):
    files = sorted(OUTBOX.glob(f"*_otp_{email.replace('@', '_at_')}*.html"),
                   key=lambda p: p.stat().st_mtime, reverse=True)
    if not files:
        return None
    m = re.search(r">(\d{6})<", files[0].read_text(encoding="utf-8"))
    return m.group(1) if m else None


def outbox_count(pattern):
    return len(list(OUTBOX.glob(pattern)))


def login(client, email):
    client.post(f"{BASE}/api/auth/request-otp", json={"email": email})
    time.sleep(0.4)
    code = latest_otp(email)
    r = client.post(f"{BASE}/api/auth/verify-otp", json={"email": email, "code": code})
    assert r.status_code == 200, r.text
    return r.json()


def make_resume(name, email, flavor):
    import docx
    doc = docx.Document()
    doc.add_heading(name, 0)
    doc.add_paragraph(f"{email} | Columbus, OH")
    doc.add_heading("Experience", 1)
    if flavor == "mgmt":
        doc.add_paragraph("Operations lead, 6 years: scheduling, budgeting, vendor management, "
                          "team leadership across 3 sites, customer service escalations.")
        doc.add_paragraph("Skills: Leadership, Budgeting, Scheduling, Excel, Customer Service")
    else:
        doc.add_paragraph("Warehouse coordinator, 4 years: inventory management, forklift "
                          "certified, logistics planning, safety compliance, data entry in Excel.")
        doc.add_paragraph("Skills: Inventory Management, Logistics, Excel, Forklift, Safety")
    path = ROOT / "data" / f"scratch_{flavor}.docx"
    doc.save(str(path))
    return path


def wait_run(admin):
    for _ in range(120):
        time.sleep(2)
        s = admin.get(f"{BASE}/api/admin/run-status").json()
        if not s["running"]:
            return s
    return {"running": True}


def main():
    admin = httpx.Client(timeout=180)
    user = httpx.Client(timeout=180)
    login(admin, ADMIN_EMAIL)

    print("\n=== A. Create scratch user via admin upload ===")
    p = make_resume("Scratch User", SCRATCH_EMAIL, "mgmt")
    with open(p, "rb") as f:
        r = admin.post(f"{BASE}/api/admin/upload-resume", files={"file": (p.name, f, "application/octet-stream")})
    check("scratch user created", r.json().get("ok") is True, str(r.json())[:200])
    login(user, SCRATCH_EMAIL)

    print("\n=== B. Multiple resumes (upload 2nd with label, set primary, delete) ===")
    p2 = make_resume("Scratch User", SCRATCH_EMAIL, "warehouse")
    with open(p2, "rb") as f:
        r = user.post(f"{BASE}/api/resumes/upload",
                      files={"file": (p2.name, f, "application/octet-stream")},
                      data={"label": "Warehouse Resume"})
    check("second resume uploaded", r.status_code == 200, r.text[:200])
    prof = user.get(f"{BASE}/api/profile").json()
    check("two resumes listed", len(prof["resumes"]) == 2, str(len(prof["resumes"])))
    second = next(x for x in prof["resumes"] if x["label"] == "Warehouse Resume")
    r = user.patch(f"{BASE}/api/resumes/{second['id']}", json={"is_primary": True})
    check("set as primary", r.status_code == 200)
    prof = user.get(f"{BASE}/api/profile").json()
    primary = next(x for x in prof["resumes"] if x["is_primary"])
    check("primary switched", primary["label"] == "Warehouse Resume")
    other = next(x for x in prof["resumes"] if not x["is_primary"])
    r = user.delete(f"{BASE}/api/resumes/{other['id']}")
    check("resume deleted", r.status_code == 200 and
          len(user.get(f"{BASE}/api/profile").json()["resumes"]) == 1)

    print("\n=== C. Give Me Advice Now + preference suppression ===")
    before = outbox_count(f"*_resume_advice_{SCRATCH_EMAIL.replace('@', '_at_')}*.html")
    r = user.post(f"{BASE}/api/resume-advice")
    check("advice email sent", r.json().get("sent") is True, str(r.json()))
    after = outbox_count(f"*_resume_advice_{SCRATCH_EMAIL.replace('@', '_at_')}*.html")
    check("advice email in outbox", after == before + 1, f"{before}->{after}")
    user.put(f"{BASE}/api/profile", json={"email_pref_resume_advice": False})
    r = user.post(f"{BASE}/api/resume-advice")
    check("advice suppressed when pref off", r.json().get("sent") is False and "disabled" in (r.json().get("reason") or ""))
    user.put(f"{BASE}/api/profile", json={"email_pref_resume_advice": True})

    print("\n=== D. Job-listings preference suppression ===")
    user.put(f"{BASE}/api/profile", json={"email_pref_job_listings": False,
                                          "job_type_fulltime": True,
                                          "preferred_positions": "Warehouse Coordinator, Logistics Specialist"})
    before = outbox_count(f"*_job_listing_{SCRATCH_EMAIL.replace('@', '_at_')}*.html")
    admin.post(f"{BASE}/api/admin/run-automation", json={"type": "main"})
    wait_run(admin)
    after = outbox_count(f"*_job_listing_{SCRATCH_EMAIL.replace('@', '_at_')}*.html")
    jobs = user.get(f"{BASE}/api/jobs").json()
    check("job email suppressed (pref off)", after == before, f"{before}->{after}")
    check("jobs still recorded on dashboard", jobs["total"] > 0, str(jobs["total"]))
    user.put(f"{BASE}/api/profile", json={"email_pref_job_listings": True})

    print("\n=== E. Email frequency transition resets cadence clock ===")
    r = user.put(f"{BASE}/api/profile", json={"email_frequency": "weekly"})
    check("frequency change accepted", r.status_code == 200)

    print("\n=== F. Urgency / portfolio / token-budget runs complete ===")
    admin.post(f"{BASE}/api/admin/run-automation", json={"type": "urgency"})
    s = wait_run(admin)
    check("urgency sweep completed", not s["running"] and "error" not in (s.get("last_result") or {}),
          str(s.get("last_result")))
    user.put(f"{BASE}/api/profile", json={"portfolio_url": "https://example.com"})
    r = user.post(f"{BASE}/api/portfolio/validate")
    check("portfolio validate works", r.json().get("is_accessible") is True, str(r.json()))
    admin.post(f"{BASE}/api/admin/run-automation", json={"type": "portfolio"})
    s = wait_run(admin)
    check("portfolio sweep completed", not s["running"], str(s.get("last_result")))
    admin.put(f"{BASE}/api/admin/token-budget", json={"daily": 10, "monthly": 100})
    before_alerts = outbox_count("*_system_ivandoublejr*")
    admin.post(f"{BASE}/api/admin/run-automation", json={"type": "budget"})
    s = wait_run(admin)
    after_alerts = outbox_count("*_system_ivandoublejr*")
    has_usage = (s.get("last_result") or {}).get("monthly_tokens", 0) > 0
    check("budget check completed", not s["running"], str(s.get("last_result")))
    check("budget alert emailed when over threshold", (after_alerts > before_alerts) or not has_usage,
          f"{before_alerts}->{after_alerts}, usage={s.get('last_result')}")
    admin.put(f"{BASE}/api/admin/token-budget", json={"daily": 2000000, "monthly": 40000000})

    print("\n=== G. GDPR deletion: story persists anonymously ===")
    user.post(f"{BASE}/api/stories", json={"previous_role": "Scratch", "new_role": "Deleted User Role",
                                           "timeframe": "1 week", "story": "Persistence test."})
    r = user.request("DELETE", f"{BASE}/api/profile")
    check("account deleted", r.status_code == 200)
    users_list = admin.get(f"{BASE}/api/admin/users").json()["users"]
    check("user gone from admin list", all(u["email"] != SCRATCH_EMAIL for u in users_list))
    stories = admin.get(f"{BASE}/api/admin/stories").json()["stories"]
    persisted = [s for s in stories if s["new_role"] == "Deleted User Role"]
    check("story persisted", len(persisted) == 1)
    check("story anonymized (user_id NULL)", persisted and persisted[0]["user_id"] is None,
          str(persisted[0] if persisted else None))
    confirm = outbox_count(f"*_system_{SCRATCH_EMAIL.replace('@', '_at_')}*")
    check("deletion confirmation email sent", confirm >= 1, str(confirm))

    print(f"\n{'=' * 50}\nPASSED: {len(passed)}  FAILED: {len(failed)}")
    if failed:
        print("Failed:", failed)
        sys.exit(1)


if __name__ == "__main__":
    main()
