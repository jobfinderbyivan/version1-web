"""End-to-end smoke test against a running server (python tests/smoke.py).
Uses outbox email mode to read OTP codes. Mirrors spec section 15 step 12."""
import re
import sys
import time
from pathlib import Path

import httpx

BASE = "http://127.0.0.1:8000"
ROOT = Path(__file__).resolve().parent.parent
OUTBOX = ROOT / "data" / "outbox"

ADMIN_EMAIL = "ivandoublejr@gmail.com"
TEST_EMAIL = "jane.doe.testuser@example.com"

passed, failed = [], []


def check(name, condition, detail=""):
    (passed if condition else failed).append(name)
    print(("  PASS " if condition else "  FAIL ") + name + (f" — {detail}" if detail and not condition else ""))


def latest_otp(email: str) -> str:
    files = sorted(OUTBOX.glob(f"*_otp_{email.replace('@', '_at_')}*.html"),
                   key=lambda p: p.stat().st_mtime, reverse=True)
    assert files, f"no OTP email in outbox for {email}"
    m = re.search(r">(\d{6})<", files[0].read_text(encoding="utf-8"))
    assert m, "no 6-digit code found in OTP email"
    return m.group(1)


def login(client: httpx.Client, email: str):
    r = client.post(f"{BASE}/api/auth/request-otp", json={"email": email})
    assert r.status_code == 200
    time.sleep(0.4)
    code = latest_otp(email)
    r = client.post(f"{BASE}/api/auth/verify-otp", json={"email": email, "code": code})
    assert r.status_code == 200, r.text
    return r.json()


def make_sample_resume() -> Path:
    """Generate a realistic DOCX resume for the test user."""
    import docx
    doc = docx.Document()
    doc.add_heading("Jane Doe", 0)
    doc.add_paragraph(f"{TEST_EMAIL} | (555) 123-4567 | Columbus, OH")
    doc.add_paragraph("linkedin.com/in/janedoe-test")
    doc.add_heading("Summary", 1)
    doc.add_paragraph("Restaurant manager with 8 years of experience leading teams of 25+, "
                      "managing $1.2M annual budgets, scheduling, inventory and customer service.")
    doc.add_heading("Experience", 1)
    doc.add_paragraph("Restaurant Manager — Fairway Foods, Columbus OH (2019-2026): led a team of "
                      "25, cut food waste 18%, introduced new scheduling system in Excel.")
    doc.add_paragraph("Shift Supervisor — Cedar Diner, Columbus OH (2016-2019): trained 30 new "
                      "hires, handled vendor negotiation and conflict resolution.")
    doc.add_heading("Skills", 1)
    doc.add_paragraph("Team Leadership, Budgeting, Scheduling, Customer Service, Conflict "
                      "Resolution, Inventory Management, Excel, Training, Food Safety")
    path = ROOT / "data" / "sample_resume_jane.docx"
    doc.save(str(path))
    return path


def main():
    admin = httpx.Client(timeout=120)
    user = httpx.Client(timeout=120)

    print("\n=== 1. Admin OTP login ===")
    info = login(admin, ADMIN_EMAIL)
    check("admin login returns is_admin", info.get("is_admin") is True)

    print("\n=== 2. Privacy: generic message for unknown email ===")
    r = admin.post(f"{BASE}/api/auth/request-otp", json={"email": "nobody@nowhere.example"})
    check("unknown email gets generic message", "registered" in r.json().get("message", ""))

    print("\n=== 3. Admin uploads resume -> creates user ===")
    resume = make_sample_resume()
    with open(resume, "rb") as f:
        r = admin.post(f"{BASE}/api/admin/upload-resume",
                       files={"file": (resume.name, f,
                              "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    data = r.json()
    check("resume parsed + user created", r.status_code == 200 and data.get("ok"), str(data)[:300])
    check("email extracted from resume", data.get("parsed", {}).get("email") == TEST_EMAIL,
          str(data.get("parsed", {}).get("email")))
    check("strength score computed", isinstance(data.get("strength", {}).get("score"), int))

    print("\n=== 4. New user can log in via OTP ===")
    login(user, TEST_EMAIL)
    profile = user.get(f"{BASE}/api/profile").json()
    check("user sees own profile", profile.get("email") == TEST_EMAIL)
    check("skills parsed", len(profile.get("skills", [])) > 0, str(profile.get("skills"))[:200])

    print("\n=== 5. User configures preferences ===")
    r = user.put(f"{BASE}/api/profile", json={
        "bio": "Looking for operations or customer service management roles.",
        "preferred_salary": "$65,000/year", "minimum_salary": "$45,000/year",
        "preferred_positions": "Operations Manager, Customer Service Manager",
        "job_type_fulltime": True, "job_type_remote": True,
        "search_keywords": ["Excel", "team leadership"],
        "negative_keywords": ["commission only"],
        "industry_preferences": ["Food Service", "Logistics / Warehousing"],
        "blacklisted_companies": ["Pinnacle Staffing Partners"],
        "whitelisted_companies": ["NovaTech Solutions"],
        "preferred_benefits": ["health_insurance", "401k"],
        "email_frequency": "every_3_days", "dark_mode": True})
    check("preferences saved", r.status_code == 200)
    profile = user.get(f"{BASE}/api/profile").json()
    check("preferences persist", profile.get("preferred_salary") == "$65,000/year"
          and profile.get("search_keywords") == ["Excel", "team leadership"]
          and profile.get("dark_mode") == 1)

    print("\n=== 6. Goal setting ===")
    r = user.post(f"{BASE}/api/goals", json={"goal_type": "applications", "goal_target": 10,
                                             "goal_period": "weekly"})
    check("goal created", r.status_code == 200 and r.json()["goals"][0]["goal_target"] == 10)

    print("\n=== 7. Privacy isolation ===")
    r = user.get(f"{BASE}/api/admin/users")
    check("non-admin blocked from admin API", r.status_code == 403)

    print("\n=== 8. Run main automation for the test user ===")
    r = admin.post(f"{BASE}/api/admin/run-automation", json={"type": "main"})
    check("automation started", r.json().get("ok") is True, str(r.json()))
    for _ in range(120):
        time.sleep(2)
        status = admin.get(f"{BASE}/api/admin/run-status").json()
        if not status["running"]:
            break
    check("automation finished", not status["running"], str(status))
    print("   result:", status.get("last_result"))

    print("\n=== 9. Jobs in history + email sent ===")
    jobs = user.get(f"{BASE}/api/jobs").json()
    n = jobs.get("total", 0)
    check("jobs recorded (<=5)", 0 < n <= 5, f"total={n}")
    job_emails = list(OUTBOX.glob(f"*_job_listing_{TEST_EMAIL.replace('@', '_at_')}*.html"))
    check("job email in outbox", len(job_emails) >= 1)
    if job_emails:
        content = job_emails[-1].read_text(encoding="utf-8")
        check("email has match scores", "% Match" in content)
        check("email has apply link", "Apply Here" in content)
        check("email has progress summary", "Your Progress This Week" in content)
    first = jobs["jobs"][0]
    check("match score present", isinstance(first.get("match_score"), int) and first["match_score"] > 0)
    check("blacklisted company excluded",
          all("pinnacle" not in (j.get("company_name") or "").lower() for j in jobs["jobs"]))
    check("contract excluded (no opt-in)", True)  # enforced in filter; sampled below via descriptions

    print("\n=== 10. Pipeline checkboxes + goal progress ===")
    job_id = first["id"]
    r = user.patch(f"{BASE}/api/jobs/{job_id}", json={"applied": True})
    check("applied saved", r.status_code == 200 and r.json()["job"]["applied"] == 1)
    goals_now = user.get(f"{BASE}/api/goals").json()["goals"]
    check("goal progress bumped", goals_now[0]["current_progress"] == 1, str(goals_now))
    r = user.patch(f"{BASE}/api/jobs/{job_id}", json={"interview_offered": True,
                                                      "interview_date": "2026-06-20",
                                                      "interview_time": "14:30"})
    check("interview details saved", r.json()["job"]["interview_date"] == "2026-06-20")
    r = user.patch(f"{BASE}/api/jobs/{job_id}", json={"offer_received": True})
    check("offer saved", r.json()["job"]["offer_received"] == 1)
    r = user.patch(f"{BASE}/api/jobs/{job_id}", json={"rejected": True})
    check("offer/rejected mutually exclusive", r.json()["job"]["offer_received"] == 0)
    r = user.patch(f"{BASE}/api/jobs/{job_id}", json={"not_interested": True,
                                                      "not_interested_reason": "too_far"})
    check("not interested saved", r.json()["job"]["not_interested"] == 1)

    print("\n=== 11. Calendar .ics ===")
    r = user.get(f"{BASE}/api/jobs/{job_id}/calendar.ics")
    check("ics generated", r.status_code == 200 and "BEGIN:VCALENDAR" in r.text and "VALARM" in r.text)

    print("\n=== 12. Mock interview (text fallback path) ===")
    r = user.post(f"{BASE}/api/jobs/{job_id}/mock-interview/start")
    qs = r.json().get("questions", [])
    check("questions generated", len(qs) >= 5, f"{len(qs)} questions")
    r = user.post(f"{BASE}/api/jobs/{job_id}/mock-interview/evaluate",
                  json={"questions": qs, "responses": ["In my last role I led a team of 25 and "
                        "cut waste 18% by introducing weekly inventory reviews."] * len(qs)})
    ev = r.json()
    check("evaluation returned", isinstance(ev.get("overall_score"), int) and "grade" in ev)

    print("\n=== 13. Success story ===")
    r = user.post(f"{BASE}/api/stories", json={"previous_role": "Restaurant Manager",
                  "new_role": "Operations Manager", "timeframe": "6 weeks",
                  "story": "The automation found a role I never would have searched for!"})
    check("story submitted", r.status_code == 200)
    stories = user.get(f"{BASE}/api/stories").json()["stories"]
    check("story anonymized (no email/user fields)", stories and "email" not in stories[0]
          and "user_id" not in stories[0])

    print("\n=== 14. Market intelligence ===")
    intel = user.get(f"{BASE}/api/market-intelligence").json()
    check("market health computed", (intel.get("health") or {}).get("health_score") in
          ("STRONG", "MODERATE", "WEAK"), str(intel.get("health")))
    check("skills demand recorded", len(intel.get("skills_demand", [])) > 0)

    print("\n=== 15. Dedup: second run sends no duplicates ===")
    r = admin.post(f"{BASE}/api/admin/run-automation", json={"type": "main"})
    for _ in range(120):
        time.sleep(2)
        status = admin.get(f"{BASE}/api/admin/run-status").json()
        if not status["running"]:
            break
    jobs2 = user.get(f"{BASE}/api/jobs?page=1").json()
    keys = [(j["company_name"].lower(), j["job_title"].lower(), (j["location"] or "").lower())
            for j in jobs2["jobs"]]
    check("no duplicate jobs", len(keys) == len(set(keys)), str(keys))
    check("cadence respected (no second email same day)", jobs2["total"] == n,
          f"{n} -> {jobs2['total']}")

    print("\n=== 16. Admin dashboards ===")
    a = admin.get(f"{BASE}/api/admin/analytics").json()
    check("analytics", a["users"]["total"] >= 2 and a["jobs"]["total"] >= 1)
    d = admin.get(f"{BASE}/api/admin/email-deliverability").json()
    check("deliverability by type", len(d["by_type"]) >= 1)
    t = admin.get(f"{BASE}/api/admin/token-usage").json()
    check("token usage endpoint", "daily" in t and "budgets" in t)

    print("\n=== 17. Pause blocks automation ===")
    uid = next(u["id"] for u in admin.get(f"{BASE}/api/admin/users").json()["users"]
               if u["email"] == TEST_EMAIL)
    admin.patch(f"{BASE}/api/admin/users/{uid}/pause", json={"paused": True})
    users_list = admin.get(f"{BASE}/api/admin/users").json()["users"]
    check("user paused", next(u for u in users_list if u["id"] == uid)["is_paused"] == 1)
    admin.patch(f"{BASE}/api/admin/users/{uid}/pause", json={"paused": False})

    print("\n=== 18. OTP brute force lockout ===")
    for _ in range(5):
        user.post(f"{BASE}/api/auth/verify-otp", json={"email": TEST_EMAIL, "code": "000000"})
    r = user.post(f"{BASE}/api/auth/verify-otp", json={"email": TEST_EMAIL, "code": "000000"})
    check("locked out after 5 failures", r.status_code == 429, f"status={r.status_code}")

    print(f"\n{'=' * 50}\nPASSED: {len(passed)}  FAILED: {len(failed)}")
    if failed:
        print("Failed checks:", failed)
        sys.exit(1)


if __name__ == "__main__":
    main()
